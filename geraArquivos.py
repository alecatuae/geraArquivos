import os
import random
import string
import json
import hashlib
import time
from io import BytesIO
from PIL import Image, ImageDraw
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd
from dataclasses import dataclass
from typing import Dict, List, Optional
from faker import Faker
from lorem_text import lorem
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import numpy as np

def carregar_configuracao(caminho_config="config.json"):
    """
    Carrega as configurações do arquivo JSON
    
    Args:
        caminho_config (str): Caminho para o arquivo de configuração JSON
        
    Returns:
        dict: Dicionário com todas as configurações carregadas
        
    Raises:
        FileNotFoundError: Se o arquivo de configuração não for encontrado
        json.JSONDecodeError: Se o arquivo JSON estiver malformado
    """
    try:
        with open(caminho_config, 'r', encoding='utf-8') as arquivo:
            return json.load(arquivo)
    except FileNotFoundError:
        print(f"⚠️  Arquivo de configuração {caminho_config} não encontrado. Usando configurações padrão.")
        return {}
    except json.JSONDecodeError as e:
        print(f"❌ Erro ao decodificar JSON: {e}")
        return {}

# Carregar configurações do arquivo JSON
CONFIG = carregar_configuracao()

# Pasta onde salvar os arquivos (carregada do config.json ou padrão)
OUTPUT_DIR = CONFIG.get("configuracao_global", {}).get("diretorio_padrao", "arquivos_teste")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Instância global do Faker para gerar dados realistas
locale_faker = CONFIG.get("configuracao_global", {}).get("locale_faker", "pt_BR")
fake = Faker(locale_faker)

@dataclass
class ConfiguracaoArquivos:
    """
    Classe de configuração para geração de arquivos de teste.
    
    Esta classe centraliza todas as configurações necessárias para gerar arquivos
    de diferentes tipos (JPEG, PDF, DOCX, XLSX, TXT) com controle sobre quantidade,
    tamanho, distribuição e parâmetros específicos por tipo de arquivo.
    
    Atributos:
        tipos_ativados (List[str]): Lista dos tipos de arquivo a serem gerados
        quantidade_por_tipo (Dict[str, int]): Quantidade específica por tipo
        quantidade_total (int): Quantidade total de arquivos (ignora quantidade_por_tipo)
        percentual_por_tipo (Dict[str, float]): Distribuição por percentual
        tamanho_mb (Dict[str, float]): Tamanho alvo em MB para cada tipo
        config_especifica (Dict[str, Dict]): Configurações específicas por tipo
        diretorio_destino (str): Diretório onde salvar os arquivos
    """
    # Tipos de arquivo ativados
    tipos_ativados: List[str] = None
    
    # Quantidade de arquivos por tipo (None = aleatório)
    quantidade_por_tipo: Dict[str, int] = None
    
    # Quantidade total de arquivos (ignora quantidade_por_tipo se especificado)
    quantidade_total: int = None
    
    # Distribuição por percentual (ex: {"pdf": 70, "outros": 30})
    percentual_por_tipo: Dict[str, float] = None
    
    # Tamanho alvo em MB para cada tipo
    tamanho_mb: Dict[str, float] = None
    
    # Configurações específicas por tipo
    config_especifica: Dict[str, Dict] = None
    
    # Diretório de destino dos arquivos
    diretorio_destino: str = None
    
    def __post_init__(self):
        """
        Inicializa valores padrão baseados no arquivo config.json.
        
        Este método é executado automaticamente após a criação da instância
        e define valores padrão para todos os atributos que não foram especificados.
        """
        # Carregar tipos de arquivo padrão do config.json
        if self.tipos_ativados is None:
            self.tipos_ativados = CONFIG.get("tipos_arquivo_padrao", ["jpeg", "pdf", "docx", "xlsx", "txt"])
        
        # Inicializar dicionário de quantidade por tipo
        if self.quantidade_por_tipo is None:
            self.quantidade_por_tipo = {}
        
        # Carregar tamanhos padrão do config.json
        if self.tamanho_mb is None:
            self.tamanho_mb = CONFIG.get("tamanhos_mb_padrao", {
                "jpeg": 0.5,    # 500KB
                "pdf": 1.0,     # 1MB
                "docx": 0.8,    # 800KB
                "xlsx": 0.3,    # 300KB
                "txt": 0.1      # 100KB
            })
        
        # Carregar configurações específicas do config.json
        if self.config_especifica is None:
            self.config_especifica = CONFIG.get("configuracoes_especificas", {
                "jpeg": {"linhas_texto": 50, "resolucao": (800, 600)},
                "pdf": {"linhas": 5, "caracteres_por_linha": 80},
                "docx": {"paragrafos": 5, "caracteres_por_paragrafo": 120},
                "xlsx": {"linhas": 20, "colunas": 15},
                "txt": {"linhas": 10, "caracteres_por_linha": 80}
            })

def obter_percentuais_padrao(template="equilibrado"):
    """
    Obtém percentuais padrão do config.json baseado em template.
    
    Esta função carrega configurações de percentual pré-definidas do config.json,
    permitindo usar templates comuns como "equilibrado", "foco_documentos", etc.
    
    Args:
        template (str): Nome do template de percentual (padrão: "equilibrado")
        
    Returns:
        Dict[str, float]: Dicionário com percentuais por tipo de arquivo
        
    Templates Disponíveis:
        - "equilibrado": Distribuição igual entre todos os tipos (20% cada)
        - "foco_documentos": Foco em PDF (40%), DOCX (30%), TXT (20%), outros (10%)
        - "foco_dados": Foco em XLSX (50%), TXT (25%), PDF (15%), outros (10%)
        - "foco_imagens": Foco em JPEG (60%), PDF (20%), outros (20%)
        - "minimal": Apenas TXT (70%) e PDF (30%)
        
    Exemplo:
        >>> percentuais = obter_percentuais_padrao("foco_documentos")
        >>> print(percentuais)
        {'pdf': 40, 'docx': 30, 'txt': 20, 'outros': 10}
    """
    percentuais_padrao = CONFIG.get("percentuais_padrao", {})
    
    if template not in percentuais_padrao:
        print(f"⚠️  Template '{template}' não encontrado. Usando 'equilibrado'.")
        template = "equilibrado"
    
    return percentuais_padrao.get(template, {
        "jpeg": 20, "pdf": 20, "docx": 20, "xlsx": 20, "txt": 20
    })

def gerar_nome_arquivo_unico(tipo_arquivo):
    """
    Gera um nome único para arquivo usando hash SHA-1.
    
    Esta função cria nomes de arquivo únicos baseados em timestamp, dados aleatórios
    e hash SHA-1, garantindo que não haja conflitos de nomes mesmo em execuções
    simultâneas ou múltiplas gerações.
    
    Args:
        tipo_arquivo (str): Extensão do arquivo (ex: "txt", "pdf", "docx")
        
    Returns:
        str: Nome único do arquivo com extensão
        
    Exemplo:
        >>> gerar_nome_arquivo_unico("txt")
        'a1b2c3d4e5f6789012345678901234567890abcd.txt'
        >>> gerar_nome_arquivo_unico("pdf")
        'f9e8d7c6b5a4938271605948372615049382716.pdf'
    """
    # Criar string única baseada em timestamp, dados aleatórios e PID
    timestamp = str(time.time())
    random_data = ''.join(random.choices(string.ascii_letters + string.digits, k=20))
    pid = str(os.getpid())
    
    # Combinar dados para criar string única
    unique_string = f"{timestamp}_{random_data}_{pid}"
    
    # Gerar hash SHA-1
    hash_object = hashlib.sha1(unique_string.encode())
    hash_hex = hash_object.hexdigest()
    
    # Retornar nome do arquivo com extensão
    return f"{hash_hex}.{tipo_arquivo}"

def texto_aleatorio(tamanho=100):
    """
    Gera uma string aleatória com caracteres alfanuméricos.
    
    Esta função é utilizada principalmente para gerar texto sobreposto em imagens JPEG
    e como conteúdo auxiliar em outros tipos de arquivo quando necessário.
    
    Args:
        tamanho (int): Número de caracteres a serem gerados (padrão: 100)
        
    Returns:
        str: String aleatória contendo letras (maiúsculas e minúsculas) e dígitos
        
    Exemplo:
        >>> texto_aleatorio(10)
        'aB3kL9mN2p'
        >>> texto_aleatorio(5)
        'Xy7Zq'
    """
    return ''.join(random.choices(string.ascii_letters + string.digits, k=tamanho))

def gerar_wordcloud_lorem(config_wordcloud, resolucao):
    """
    Gera um wordcloud com palavras Lorem Ipsum e frequências aleatórias.
    
    Esta função cria um wordcloud visualmente atrativo usando palavras do Lorem Ipsum
    com frequências variadas, cores aleatórias e configurações personalizáveis.
    O resultado é uma imagem colorida e dinâmica ideal para arquivos JPEG e PNG.
    
    Args:
        config_wordcloud (dict): Configurações específicas do wordcloud
            - max_palavras: Número máximo de palavras a incluir
            - largura, altura: Dimensões do wordcloud
            - background_color: Cor de fundo
            - colormap: Mapa de cores (viridis, plasma, etc.)
            - max_font_size, min_font_size: Tamanhos de fonte
            - relative_scaling: Escala relativa das palavras
            - prefer_horizontal: Preferência por orientação horizontal
        resolucao (tuple): Resolução final (largura, altura)
        
    Returns:
        PIL.Image: Imagem do wordcloud como objeto PIL
        
    Características:
        - Palavras Lorem Ipsum clássicas
        - Frequências variadas e realistas
        - Cores vibrantes e aleatórias
        - Layout otimizado para legibilidade
        - Suporte a transparência (PNG)
        
    Exemplo:
        >>> config = {"max_palavras": 50, "colormap": "viridis"}
        >>> img = gerar_wordcloud_lorem(config, (800, 600))
        >>> img.save("wordcloud.png")
    """
    # Carregar configurações de wordcloud do config.json
    config_global = CONFIG.get("configuracoes_wordcloud", {})
    palavras_lorem = config_global.get("palavras_lorem", [
        "lorem", "ipsum", "dolor", "sit", "amet", "consectetur", "adipiscing", "elit",
        "sed", "do", "eiusmod", "tempor", "incididunt", "ut", "labore", "et", "dolore"
    ])
    frequencias_padrao = config_global.get("frequencias_padrao", {
        "lorem": 10, "ipsum": 8, "dolor": 6, "sit": 5, "amet": 4
    })
    cores_disponiveis = config_global.get("cores_disponiveis", ["viridis", "plasma", "inferno"])
    
    # Configurações do wordcloud
    max_palavras = config_wordcloud.get("max_palavras", 100)
    largura = config_wordcloud.get("largura", 800)
    altura = config_wordcloud.get("altura", 600)
    background_color = config_wordcloud.get("background_color", "white")
    colormap = config_wordcloud.get("colormap", "viridis")
    max_font_size = config_wordcloud.get("max_font_size", 100)
    min_font_size = config_wordcloud.get("min_font_size", 10)
    relative_scaling = config_wordcloud.get("relative_scaling", 0.5)
    prefer_horizontal = config_wordcloud.get("prefer_horizontal", 0.9)
    
    # Selecionar palavras aleatórias do Lorem Ipsum
    palavras_selecionadas = random.sample(palavras_lorem, min(max_palavras, len(palavras_lorem)))
    
    # Gerar frequências variadas para as palavras
    frequencias = {}
    for palavra in palavras_selecionadas:
        if palavra in frequencias_padrao:
            # Usar frequência padrão com variação aleatória
            base_freq = frequencias_padrao[palavra]
            variacao = random.uniform(0.5, 2.0)
            frequencias[palavra] = int(base_freq * variacao)
        else:
            # Frequência aleatória para palavras não padrão
            frequencias[palavra] = random.randint(1, 5)
    
    # Garantir que pelo menos algumas palavras tenham frequência alta
    palavras_principais = random.sample(palavras_selecionadas, min(5, len(palavras_selecionadas)))
    for palavra in palavras_principais:
        frequencias[palavra] = random.randint(8, 15)
    
    # Criar wordcloud
    wordcloud = WordCloud(
        width=largura,
        height=altura,
        background_color=background_color,
        colormap=colormap,
        max_font_size=max_font_size,
        min_font_size=min_font_size,
        relative_scaling=relative_scaling,
        prefer_horizontal=prefer_horizontal,
        max_words=max_palavras,
        random_state=42  # Para reprodutibilidade
    ).generate_from_frequencies(frequencias)
    
    # Converter para PIL Image
    img_array = wordcloud.to_array()
    img_pil = Image.fromarray(img_array)
    
    # Redimensionar se necessário
    if resolucao != (largura, altura):
        img_pil = img_pil.resize(resolucao, Image.Resampling.LANCZOS)
    
    return img_pil

def gerar_dados_realistas_xlsx(num_linhas):
    """
    Gera dados realistas para planilhas XLSX usando a biblioteca Faker.
    
    Esta função cria um conjunto completo de dados realistas em português brasileiro
    para simular uma planilha de funcionários ou clientes. Os dados incluem informações
    pessoais, profissionais e de contato, todos gerados de forma consistente e realista.
    
    Args:
        num_linhas (int): Número de linhas (registros) a serem gerados
        
    Returns:
        Dict[str, List]: Dicionário onde cada chave é o nome da coluna e cada valor
                        é uma lista com os dados para aquela coluna
        
    Colunas Geradas:
        - ID: Números únicos de 1000-9999
        - Nome: Nomes completos brasileiros
        - Email: Endereços de email realistas
        - Telefone: Números de telefone no formato brasileiro
        - Endereço: Endereços completos (bairro, cidade, estado)
        - Cidade: Nomes de cidades brasileiras
        - Estado: Estados brasileiros
        - CEP: Códigos postais válidos
        - Data_Nascimento: Datas de nascimento (18-80 anos)
        - Profissão: Profissões diversas
        - Empresa: Nomes de empresas
        - Salário: Valores salariais (R$ 1.500 - R$ 15.000)
        - Data_Contrato: Datas de contratação (últimos 5 anos)
        - Status: Status do funcionário (Ativo, Inativo, Férias, Licença)
        - Observações: Textos aleatórios de até 100 caracteres
        
    Exemplo:
        >>> dados = gerar_dados_realistas_xlsx(3)
        >>> print(dados['Nome'][:2])
        ['João Silva Santos', 'Maria Oliveira Costa']
        >>> print(dados['Email'][:2])
        ['joao.silva@email.com', 'maria.oliveira@empresa.com.br']
    """
    # Carregar configurações do Faker do config.json
    config_faker = CONFIG.get("configuracoes_faker", {})
    id_min = config_faker.get("id_minimo", 1000)
    id_max = config_faker.get("id_maximo", 9999)
    idade_min = config_faker.get("idade_minima", 18)
    idade_max = config_faker.get("idade_maxima", 80)
    salario_min = config_faker.get("salario_minimo", 1500)
    salario_max = config_faker.get("salario_maximo", 15000)
    anos_contrato = config_faker.get("anos_contrato", 5)
    status_opcoes = config_faker.get("status_funcionario", ["Ativo", "Inativo", "Férias", "Licença"])
    tamanho_obs = config_faker.get("tamanho_observacoes", 100)
    
    dados = {
        "ID": [fake.random_int(min=id_min, max=id_max) for _ in range(num_linhas)],
        "Nome": [fake.name() for _ in range(num_linhas)],
        "Email": [fake.email() for _ in range(num_linhas)],
        "Telefone": [fake.phone_number() for _ in range(num_linhas)],
        "Endereço": [fake.address().replace('\n', ', ') for _ in range(num_linhas)],
        "Cidade": [fake.city() for _ in range(num_linhas)],
        "Estado": [fake.state() for _ in range(num_linhas)],
        "CEP": [fake.postcode() for _ in range(num_linhas)],
        "Data_Nascimento": [fake.date_of_birth(minimum_age=idade_min, maximum_age=idade_max).strftime('%d/%m/%Y') for _ in range(num_linhas)],
        "Profissão": [fake.job() for _ in range(num_linhas)],
        "Empresa": [fake.company() for _ in range(num_linhas)],
        "Salário": [fake.random_int(min=salario_min, max=salario_max) for _ in range(num_linhas)],
        "Data_Contrato": [fake.date_between(start_date=f'-{anos_contrato}y', end_date='today').strftime('%d/%m/%Y') for _ in range(num_linhas)],
        "Status": [fake.random_element(elements=status_opcoes) for _ in range(num_linhas)],
        "Observações": [fake.text(max_nb_chars=tamanho_obs) for _ in range(num_linhas)]
    }
    return dados

def gerar_texto_lorem_ipsum(tamanho_mb_alvo, tipo_arquivo="txt"):
    """
    Gera texto Lorem Ipsum baseado no tamanho alvo em MB.
    
    Esta função utiliza a biblioteca lorem-text para gerar texto Lorem Ipsum
    clássico, ajustando automaticamente a quantidade de texto baseada no tamanho
    desejado em MB. Considera o overhead de formatação de diferentes tipos de arquivo.
    
    Args:
        tamanho_mb_alvo (float): Tamanho alvo em MB para o arquivo
        tipo_arquivo (str): Tipo do arquivo ("txt", "pdf", "docx") para ajustar overhead
        
    Returns:
        List[str]: Lista de strings contendo parágrafos de Lorem Ipsum
        
    Comportamento por Tipo de Arquivo:
        - TXT: 1 caractere ≈ 1 byte (controle preciso)
        - PDF: Redução de 30% devido ao overhead de formatação
        - DOCX: Redução de 50% devido ao overhead de XML
        
    Exemplo:
        >>> textos = gerar_texto_lorem_ipsum(0.1, "txt")
        >>> print(len(textos))
        3
        >>> print(textos[0][:50])
        Lorem ipsum dolor sit amet, consectetur adipiscing elit...
    """
    # Carregar configurações do Lorem Ipsum do config.json
    config_lorem = CONFIG.get("configuracoes_lorem_ipsum", {})
    tamanho_min = config_lorem.get("tamanho_paragrafo_minimo", 50)
    tamanho_max = config_lorem.get("tamanho_paragrafo_maximo", 200)
    reducao_pdf = config_lorem.get("reducao_pdf", 0.7)
    reducao_docx = config_lorem.get("reducao_docx", 0.5)
    caracteres_min = config_lorem.get("caracteres_minimos", 1000)
    
    # Calcular quantidade de caracteres necessários baseado no tamanho
    # Assumindo ~1 caractere = 1 byte para texto simples
    caracteres_necessarios = int(tamanho_mb_alvo * 1024 * 1024)
    
    # Ajustar baseado no tipo de arquivo
    if tipo_arquivo == "pdf":
        # PDF tem overhead de formatação, reduzir conforme configurado
        caracteres_necessarios = int(caracteres_necessarios * reducao_pdf)
    elif tipo_arquivo == "docx":
        # DOCX tem overhead de XML, reduzir conforme configurado
        caracteres_necessarios = int(caracteres_necessarios * reducao_docx)
    
    # Garantir mínimo de caracteres
    caracteres_necessarios = max(caracteres_necessarios, caracteres_min)
    
    textos = []
    caracteres_gerados = 0
    
    while caracteres_gerados < caracteres_necessarios:
        # Gerar parágrafo com tamanho variável
        tamanho_paragrafo = random.randint(tamanho_min, tamanho_max)
        paragrafo = lorem.paragraph()
        
        # Se o parágrafo for muito pequeno, gerar mais texto
        while len(paragrafo) < 100:
            paragrafo += " " + lorem.sentence()
        
        textos.append(paragrafo)
        caracteres_gerados += len(paragrafo)
        
        # Adicionar quebra de linha para TXT
        if tipo_arquivo == "txt":
            caracteres_gerados += 1  # \n
    
    return textos

def gerar_texto_lorem_por_linhas(num_linhas, caracteres_por_linha=80):
    """
    Gera texto Lorem Ipsum com número específico de linhas.
    
    Esta função gera texto Lorem Ipsum organizado em linhas, onde cada linha
    tem um tamanho controlado. É útil para gerar arquivos TXT com estrutura
    específica ou quando se deseja controle preciso sobre o número de linhas.
    
    Args:
        num_linhas (int): Número de linhas a serem geradas
        caracteres_por_linha (int): Número máximo de caracteres por linha (padrão: 80)
        
    Returns:
        List[str]: Lista de strings, cada uma representando uma linha de texto
        
    Características:
        - Cada linha tem tamanho variável (50% a 100% do tamanho máximo)
        - Quebra inteligente de palavras para não cortar no meio
        - Numeração automática para linhas curtas (≤100 caracteres)
        - Texto Lorem Ipsum clássico e profissional
        
    Exemplo:
        >>> linhas = gerar_texto_lorem_por_linhas(3, 50)
        >>> print(linhas[0])
        Linha 1: Lorem ipsum dolor sit amet consectetur adipiscing
        >>> print(len(linhas))
        3
    """
    # Carregar configurações do Lorem Ipsum do config.json
    config_lorem = CONFIG.get("configuracoes_lorem_ipsum", {})
    tamanho_linha_min = config_lorem.get("tamanho_linha_minimo", 50)
    tamanho_linha_max = config_lorem.get("tamanho_linha_maximo", 200)
    
    linhas = []
    
    for i in range(num_linhas):
        # Gerar linha com tamanho variável
        tamanho_linha = random.randint(caracteres_por_linha // 2, caracteres_por_linha)
        
        # Gerar texto até atingir o tamanho desejado
        linha = ""
        while len(linha) < tamanho_linha:
            palavra = lorem.word()
            if len(linha) + len(palavra) + 1 <= tamanho_linha:
                linha += palavra + " " if linha else palavra
            else:
                break
        
        # Adicionar número da linha se for TXT (linhas curtas)
        if caracteres_por_linha <= 100:  # Assumindo que é TXT
            linha = f"Linha {i+1}: {linha.strip()}"
        
        linhas.append(linha.strip())
    
    return linhas

def calcular_tamanho_arquivo(caminho_arquivo):
    """
    Calcula o tamanho de um arquivo em MB.
    
    Esta função utilitária é usada para verificar o tamanho real dos arquivos
    gerados e comparar com o tamanho alvo configurado. É útil para validação
    e feedback durante o processo de geração.
    
    Args:
        caminho_arquivo (str): Caminho completo para o arquivo
        
    Returns:
        float: Tamanho do arquivo em MB (com precisão de 2 casas decimais)
        
    Raises:
        FileNotFoundError: Se o arquivo não existir
        OSError: Se houver erro ao acessar o arquivo
        
    Exemplo:
        >>> tamanho = calcular_tamanho_arquivo("arquivo.txt")
        >>> print(f"Tamanho: {tamanho:.2f} MB")
        Tamanho: 0.15 MB
    """
    return os.path.getsize(caminho_arquivo) / (1024 * 1024)

# Funções auxiliares para distribuição por percentual
def calcular_distribuicao_por_percentual(quantidade_total, percentual_por_tipo, tipos_ativados):
    """
    Calcula a distribuição de arquivos baseada em percentuais
    
    Args:
        quantidade_total: Quantidade total de arquivos
        percentual_por_tipo: Dicionário com percentuais por tipo
        tipos_ativados: Lista de tipos ativados
    
    Returns:
        Dict com quantidade por tipo
    """
    if not percentual_por_tipo:
        return {}
    
    # Verificar se a soma dos percentuais é 100%
    soma_percentuais = sum(percentual_por_tipo.values())
    if abs(soma_percentuais - 100.0) > 0.01:  # Tolerância para arredondamento
        print(f"⚠️  Aviso: Soma dos percentuais ({soma_percentuais}%) não é 100%")
    
    distribuicao = {}
    total_distribuido = 0
    
    # Distribuir arquivos baseado nos percentuais
    for tipo, percentual in percentual_por_tipo.items():
        if tipo == "outros":
            # Distribuir percentual "outros" entre tipos restantes
            tipos_restantes = [t for t in tipos_ativados if t not in distribuicao]
            if tipos_restantes:
                percentual_por_tipo_restante = percentual / len(tipos_restantes)
                for tipo_restante in tipos_restantes:
                    quantidade = int(quantidade_total * percentual_por_tipo_restante / 100)
                    distribuicao[tipo_restante] = quantidade
                    total_distribuido += quantidade
        else:
            quantidade = int(quantidade_total * percentual / 100)
            distribuicao[tipo] = quantidade
            total_distribuido += quantidade
    
    # Ajustar diferenças de arredondamento
    diferenca = quantidade_total - total_distribuido
    if diferenca != 0:
        # Adicionar a diferença ao tipo com maior quantidade
        tipo_maior = max(distribuicao.keys(), key=lambda k: distribuicao[k])
        distribuicao[tipo_maior] += diferenca
    
    return distribuicao

def validar_percentuais(percentual_por_tipo):
    """
    Valida se os percentuais estão corretos
    
    Args:
        percentual_por_tipo: Dicionário com percentuais
    
    Returns:
        bool: True se válido, False caso contrário
    """
    if not percentual_por_tipo:
        return True
    
    soma = sum(percentual_por_tipo.values())
    return 99.0 <= soma <= 101.0  # Tolerância para arredondamento

def ajustar_conteudo_para_tamanho(tipo_arquivo, tamanho_mb_alvo, config):
    """Ajusta o conteúdo baseado no tamanho alvo"""
    if tipo_arquivo == "txt":
        # Para TXT: usar Lorem Ipsum baseado no tamanho
        return tamanho_mb_alvo
    
    elif tipo_arquivo == "pdf":
        # Para PDF: usar Lorem Ipsum baseado no tamanho
        return tamanho_mb_alvo
    
    elif tipo_arquivo == "docx":
        # Para DOCX: usar Lorem Ipsum baseado no tamanho
        return tamanho_mb_alvo
    
    elif tipo_arquivo == "xlsx":
        # Para XLSX: estimativa baseada em linhas
        # Considerando ~15 colunas com dados realistas, cada linha ≈ 0.5KB
        # Então 1MB ≈ 2000 linhas
        linhas_necessarias = int(tamanho_mb_alvo * 2000)
        return max(linhas_necessarias, 10)  # Mínimo de 10 linhas
    
    elif tipo_arquivo == "jpeg":
        # Para JPEG: ajustar resolução baseada no tamanho
        if tamanho_mb_alvo <= 0.1:
            return (400, 300)
        elif tamanho_mb_alvo <= 0.5:
            return (800, 600)
        elif tamanho_mb_alvo <= 1.0:
            return (1200, 900)
        else:
            return (1600, 1200)
    
    elif tipo_arquivo == "png":
        # Para PNG: ajustar resolução baseada no tamanho (PNG é maior que JPEG)
        if tamanho_mb_alvo <= 0.1:
            return (400, 300)
        elif tamanho_mb_alvo <= 0.5:
            return (800, 600)
        elif tamanho_mb_alvo <= 1.0:
            return (1024, 768)
        elif tamanho_mb_alvo <= 1.5:
            return (1280, 960)
        else:
            return (1600, 1200)
    
    return config

def gerar_jpeg(nome, config, tamanho_mb_alvo=None):
    """
    Gera um arquivo JPEG com wordcloud Lorem Ipsum.
    
    Esta função cria uma imagem JPEG com um wordcloud colorido e dinâmico
    usando palavras do Lorem Ipsum com frequências variadas. A resolução
    é ajustada automaticamente baseada no tamanho alvo em MB.
    
    Args:
        nome (str): Caminho completo onde salvar o arquivo JPEG
        config (dict): Configurações específicas para JPEG
            - resolucao: Tupla (largura, altura) da imagem
            - wordcloud: Configurações do wordcloud
        tamanho_mb_alvo (float, optional): Tamanho alvo em MB para ajustar resolução
        
    Características:
        - Wordcloud com palavras Lorem Ipsum
        - Cores vibrantes e aleatórias
        - Frequências variadas para realismo
        - Resolução ajustada automaticamente
        - Qualidade JPEG configurável
        
    Exemplo:
        >>> config = {"resolucao": (800, 600), "wordcloud": {"max_palavras": 50}}
        >>> gerar_jpeg("imagem.jpg", config, 0.5)
        # Gera JPEG com wordcloud de ~0.5MB
    """
    # Carregar configurações específicas do JPEG do config.json
    config_jpeg = CONFIG.get("configuracoes_especificas", {}).get("jpeg", {})
    qualidade = config_jpeg.get("qualidade", 85)
    formato_cor = config_jpeg.get("formato_cor", "RGB")
    config_wordcloud = config_jpeg.get("wordcloud", {})
    
    resolucao = config["resolucao"]
    if tamanho_mb_alvo:
        resolucao = ajustar_conteudo_para_tamanho("jpeg", tamanho_mb_alvo, config)
    
    # Gerar wordcloud Lorem Ipsum
    img = gerar_wordcloud_lorem(config_wordcloud, resolucao)
    
    # Converter para RGB se necessário (JPEG não suporta transparência)
    if img.mode != "RGB":
        # Criar fundo branco para transparência
        fundo = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "RGBA":
            fundo.paste(img, mask=img.split()[-1])  # Usar canal alpha como máscara
        else:
            fundo.paste(img)
        img = fundo
    
    # Salvar com qualidade configurável
    img.save(nome, "JPEG", quality=qualidade)

def gerar_png(nome, config, tamanho_mb_alvo=None):
    """
    Gera um arquivo PNG com wordcloud Lorem Ipsum e transparência.
    
    Esta função cria uma imagem PNG com um wordcloud colorido e dinâmico
    usando palavras do Lorem Ipsum, com suporte a transparência. A resolução
    é ajustada automaticamente baseada no tamanho alvo em MB.
    
    Args:
        nome (str): Caminho completo onde salvar o arquivo PNG
        config (dict): Configurações específicas para PNG
            - resolucao: Tupla (largura, altura) da imagem
            - wordcloud: Configurações do wordcloud
            - incluir_transparencia: Se deve incluir transparência
        tamanho_mb_alvo (float, optional): Tamanho alvo em MB para ajustar resolução
        
    Características:
        - Wordcloud com palavras Lorem Ipsum
        - Cores vibrantes e aleatórias
        - Suporte a transparência (canal alpha)
        - Frequências variadas para realismo
        - Resolução ajustada automaticamente
        - Compressão PNG configurável
        
    Exemplo:
        >>> config = {"resolucao": (1024, 768), "wordcloud": {"max_palavras": 60}}
        >>> gerar_png("imagem.png", config, 0.6)
        # Gera PNG com wordcloud de ~0.6MB
    """
    # Carregar configurações específicas do PNG do config.json
    config_png = CONFIG.get("configuracoes_especificas", {}).get("png", {})
    formato_cor = config_png.get("formato_cor", "RGBA")
    incluir_transparencia = config_png.get("incluir_transparencia", True)
    compressao = config_png.get("compressao", 6)
    config_wordcloud = config_png.get("wordcloud", {})
    
    resolucao = config["resolucao"]
    if tamanho_mb_alvo:
        resolucao = ajustar_conteudo_para_tamanho("png", tamanho_mb_alvo, config)
    
    # Gerar wordcloud Lorem Ipsum
    img = gerar_wordcloud_lorem(config_wordcloud, resolucao)
    
    # Ajustar transparência se necessário
    if incluir_transparencia and img.mode != "RGBA":
        # Converter para RGBA se não for transparente
        if img.mode == "RGB":
            img = img.convert("RGBA")
        else:
            img = img.convert("RGBA")
    elif not incluir_transparencia and img.mode == "RGBA":
        # Converter para RGB se não deve ter transparência
        fundo = Image.new("RGB", img.size, (255, 255, 255))
        fundo.paste(img, mask=img.split()[-1])
        img = fundo
    
    # Salvar com compressão configurável
    img.save(nome, "PNG", compress_level=compressao)

def gerar_pdf(nome, config, tamanho_mb_alvo=None):
    """
    Gera um arquivo PDF com texto Lorem Ipsum formatado.
    
    Esta função cria um documento PDF profissional com texto Lorem Ipsum,
    incluindo quebra de página automática, formatação adequada e controle
    de tamanho baseado em MB. O texto é quebrado inteligentemente em linhas
    respeitando o limite de caracteres configurado.
    
    Args:
        nome (str): Caminho completo onde salvar o arquivo PDF
        config (dict): Configurações específicas para PDF
            - linhas: Número de linhas (quando não usando tamanho alvo)
            - caracteres_por_linha: Limite de caracteres por linha
        tamanho_mb_alvo (float, optional): Tamanho alvo em MB para ajustar conteúdo
        
    Características:
        - Texto Lorem Ipsum clássico e profissional
        - Quebra de página automática quando necessário
        - Quebra de linha inteligente (não corta palavras)
        - Margens e espaçamento configuráveis
        - Controle de tamanho baseado em MB
        
    Exemplo:
        >>> config = {"linhas": 10, "caracteres_por_linha": 80}
        >>> gerar_pdf("documento.pdf", config, 0.5)
        # Gera PDF de ~0.5MB
    """
    # Carregar configurações específicas do PDF do config.json
    config_pdf = CONFIG.get("configuracoes_especificas", {}).get("pdf", {})
    margem_esq = config_pdf.get("margem_esquerda", 100)
    margem_sup = config_pdf.get("margem_superior", 800)
    espacamento = config_pdf.get("espacamento_linhas", 20)
    altura_min = config_pdf.get("altura_minima_pagina", 50)
    fonte_tam = config_pdf.get("fonte_tamanho", 12)
    
    if tamanho_mb_alvo:
        # Usar Lorem Ipsum baseado no tamanho
        textos = gerar_texto_lorem_ipsum(tamanho_mb_alvo, "pdf")
    else:
        # Usar configuração padrão
        linhas = config["linhas"]
        textos = gerar_texto_lorem_por_linhas(linhas, config["caracteres_por_linha"])
    
    c = canvas.Canvas(nome)
    y_pos = margem_sup
    
    for texto in textos:
        # Quebrar texto em linhas se necessário
        palavras = texto.split()
        linha_atual = ""
        
        for palavra in palavras:
            if len(linha_atual + " " + palavra) <= config["caracteres_por_linha"]:
                linha_atual += " " + palavra if linha_atual else palavra
            else:
                if linha_atual:
                    c.drawString(margem_esq, y_pos, linha_atual)
                    y_pos -= espacamento
                    if y_pos < altura_min:  # Nova página se necessário
                        c.showPage()
                        y_pos = margem_sup
                linha_atual = palavra
        
        # Desenhar última linha
        if linha_atual:
            c.drawString(margem_esq, y_pos, linha_atual)
            y_pos -= espacamento
            if y_pos < altura_min:
                c.showPage()
                y_pos = margem_sup
    
    c.save()

# Gerar DOCX
def gerar_docx(nome, config, tamanho_mb_alvo=None):
    if tamanho_mb_alvo:
        # Usar Lorem Ipsum baseado no tamanho
        textos = gerar_texto_lorem_ipsum(tamanho_mb_alvo, "docx")
    else:
        # Usar configuração padrão
        paragrafos = config["paragrafos"]
        textos = []
        for i in range(paragrafos):
            # Gerar parágrafo Lorem Ipsum
            paragrafo = lorem.paragraph()
            while len(paragrafo) < config["caracteres_por_paragrafo"]:
                paragrafo += " " + lorem.sentence()
            textos.append(paragrafo)
    
    doc = Document()
    
    # Adicionar título
    doc.add_heading('Documento Lorem Ipsum', 0)
    
    # Adicionar parágrafos
    for texto in textos:
        doc.add_paragraph(texto)
    
    # Adicionar informações do documento
    doc.add_heading('Informações do Documento', level=1)
    doc.add_paragraph(f'Data de geração: {fake.date_time_this_year().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'ID do documento: {fake.uuid4()}')
    doc.add_paragraph(f'Tamanho alvo: {tamanho_mb_alvo:.2f} MB' if tamanho_mb_alvo else 'Tamanho padrão')
    
    doc.save(nome)

# Gerar XLSX
def gerar_xlsx(nome, config, tamanho_mb_alvo=None):
    linhas = config["linhas"]
    if tamanho_mb_alvo:
        linhas = ajustar_conteudo_para_tamanho("xlsx", tamanho_mb_alvo, config)
    
    # Gerar dados realistas usando Faker
    dados = gerar_dados_realistas_xlsx(linhas)
    df = pd.DataFrame(dados)
    
    # Salvar com formatação melhorada
    with pd.ExcelWriter(nome, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Dados', index=False)
        
        # Ajustar largura das colunas automaticamente
        worksheet = writer.sheets['Dados']
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)  # Máximo de 50 caracteres
            worksheet.column_dimensions[column_letter].width = adjusted_width

# Gerar TXT
def gerar_txt(nome, config, tamanho_mb_alvo=None):
    if tamanho_mb_alvo:
        # Usar Lorem Ipsum baseado no tamanho
        textos = gerar_texto_lorem_ipsum(tamanho_mb_alvo, "txt")
    else:
        # Usar configuração padrão
        linhas = config["linhas"]
        textos = gerar_texto_lorem_por_linhas(linhas, config["caracteres_por_linha"])
    
    with open(nome, 'w', encoding='utf-8') as arquivo:
        # Escrever cabeçalho
        arquivo.write("=" * 80 + "\n")
        arquivo.write("DOCUMENTO LOREM IPSUM\n")
        arquivo.write("=" * 80 + "\n\n")
        
        # Escrever conteúdo Lorem Ipsum
        for i, texto in enumerate(textos, 1):
            arquivo.write(f"Parágrafo {i}:\n")
            arquivo.write(texto + "\n\n")
        
        # Escrever rodapé
        arquivo.write("=" * 80 + "\n")
        arquivo.write("INFORMAÇÕES DO DOCUMENTO\n")
        arquivo.write("=" * 80 + "\n")
        arquivo.write(f"Data de geração: {fake.date_time_this_year().strftime('%d/%m/%Y %H:%M')}\n")
        arquivo.write(f"ID do arquivo: {fake.uuid4()}\n")
        arquivo.write(f"Tamanho alvo: {tamanho_mb_alvo:.2f} MB\n" if tamanho_mb_alvo else "Tamanho padrão\n")
        arquivo.write(f"Total de parágrafos: {len(textos)}\n")
        arquivo.write("=" * 80 + "\n")

def gerar_arquivos(config: ConfiguracaoArquivos = None, qtd_total=None):
    """
    Função principal para geração de arquivos de teste.
    
    Esta é a função central do sistema que coordena a geração de todos os tipos
    de arquivo (JPEG, PDF, DOCX, XLSX, TXT) baseada na configuração fornecida.
    Suporta múltiplos modos de operação: quantidade específica por tipo,
    distribuição aleatória, distribuição por percentual e controle de tamanho.
    
    Args:
        config (ConfiguracaoArquivos, optional): Configuração completa para geração.
            Se None, usa configurações padrão do config.json
        qtd_total (int, optional): Quantidade total de arquivos (ignora quantidade_por_tipo)
        
    Modos de Operação:
        1. **Quantidade por Tipo**: Especifica exatamente quantos arquivos de cada tipo
        2. **Quantidade Total**: Gera quantidade total distribuída aleatoriamente
        3. **Distribuição por Percentual**: Controla percentual de cada tipo
        4. **Modo Aleatório**: Distribui arquivos aleatoriamente entre tipos ativados
        
    Características:
        - Criação automática de diretórios
        - Controle de tamanho por tipo de arquivo
        - Geração de dados realistas (Faker + Lorem Ipsum)
        - Feedback detalhado do progresso
        - Tratamento de erros robusto
        
    Exemplo:
        >>> # Configuração básica
        >>> config = ConfiguracaoArquivos(
        ...     tipos_ativados=["txt", "pdf"],
        ...     quantidade_por_tipo={"txt": 3, "pdf": 2}
        ... )
        >>> gerar_arquivos(config)
        
        >>> # Modo aleatório
        >>> gerar_arquivos(qtd_total=10)
        
        >>> # Com diretório personalizado
        >>> config.diretorio_destino = "meus_arquivos"
        >>> gerar_arquivos(config)
    """
    if config is None:
        config = ConfiguracaoArquivos()
    
    # Determinar diretório de destino
    diretorio_destino = config.diretorio_destino if config.diretorio_destino else OUTPUT_DIR
    os.makedirs(diretorio_destino, exist_ok=True)
    
    # Determinar quantos arquivos gerar de cada tipo
    arquivos_para_gerar = {}
    
    # Prioridade: quantidade_total > qtd_total > quantidade_por_tipo > aleatório
    if config.quantidade_total is not None:
        qtd_total = config.quantidade_total
    
    if qtd_total is not None:
        # Verificar se há distribuição por percentual
        if config.percentual_por_tipo:
            # Validar percentuais
            if not validar_percentuais(config.percentual_por_tipo):
                print("❌ Erro: Percentuais inválidos (soma deve ser 100%)")
                return
            
            # Calcular distribuição por percentual
            arquivos_para_gerar = calcular_distribuicao_por_percentual(
                qtd_total, config.percentual_por_tipo, config.tipos_ativados
            )
            
            # Mostrar distribuição calculada
            print(f"📊 Distribuição por percentual:")
            for tipo, qtd in arquivos_para_gerar.items():
                percentual = (qtd / qtd_total) * 100
                print(f"   {tipo.upper()}: {qtd} arquivos ({percentual:.1f}%)")
            print()
        else:
            # Modo aleatório: distribuir qtd_total entre tipos ativados
            for i in range(qtd_total):
                tipo = random.choice(config.tipos_ativados)
                arquivos_para_gerar[tipo] = arquivos_para_gerar.get(tipo, 0) + 1
    else:
        # Modo controlado: usar quantidade_por_tipo
        for tipo in config.tipos_ativados:
            if tipo in config.quantidade_por_tipo:
                arquivos_para_gerar[tipo] = config.quantidade_por_tipo[tipo]
            else:
                # Se não especificado, gerar 1 arquivo do tipo
                arquivos_para_gerar[tipo] = 1
    
    # Gerar os arquivos
    total_gerado = 0
    
    for tipo, quantidade in arquivos_para_gerar.items():
        for i in range(quantidade):
            # Gerar nome único usando SHA-1
            nome_arquivo = gerar_nome_arquivo_unico(tipo)
            nome = os.path.join(diretorio_destino, nome_arquivo)
            tamanho_alvo = config.tamanho_mb.get(tipo, 0.5)
            config_tipo = config.config_especifica.get(tipo, {})
            
            try:
                if tipo == "jpeg":
                    gerar_jpeg(nome, config_tipo, tamanho_alvo)
                elif tipo == "png":
                    gerar_png(nome, config_tipo, tamanho_alvo)
                elif tipo == "pdf":
                    gerar_pdf(nome, config_tipo, tamanho_alvo)
                elif tipo == "docx":
                    gerar_docx(nome, config_tipo, tamanho_alvo)
                elif tipo == "xlsx":
                    gerar_xlsx(nome, config_tipo, tamanho_alvo)
                elif tipo == "txt":
                    gerar_txt(nome, config_tipo, tamanho_alvo)
                
                tamanho_real = calcular_tamanho_arquivo(nome)
                print(f"[OK] Gerado: {nome} ({tamanho_real:.2f} MB)")
                total_gerado += 1
                
            except Exception as e:
                print(f"[ERRO] Falha ao gerar {nome}: {e}")
    
    print(f"\n✅ Total de arquivos gerados: {total_gerado}")

# Funções de conveniência para configurações comuns
def gerar_arquivos_aleatorios(qtd=20, tipos_ativados=None, diretorio_destino=None):
    """Gera arquivos aleatoriamente"""
    config = ConfiguracaoArquivos()
    if tipos_ativados:
        config.tipos_ativados = tipos_ativados
    if diretorio_destino:
        config.diretorio_destino = diretorio_destino
    gerar_arquivos(config, qtd)

def gerar_arquivos_por_tipo(quantidade_por_tipo, tamanhos_mb=None, diretorio_destino=None):
    """Gera arquivos com quantidade específica por tipo"""
    config = ConfiguracaoArquivos()
    config.quantidade_por_tipo = quantidade_por_tipo
    if tamanhos_mb:
        config.tamanho_mb.update(tamanhos_mb)
    if diretorio_destino:
        config.diretorio_destino = diretorio_destino
    gerar_arquivos(config)

def gerar_arquivos_por_quantidade(quantidade_total, tipos_ativados=None, diretorio_destino=None):
    """Gera arquivos com quantidade total específica"""
    config = ConfiguracaoArquivos()
    config.quantidade_total = quantidade_total
    if tipos_ativados:
        config.tipos_ativados = tipos_ativados
    if diretorio_destino:
        config.diretorio_destino = diretorio_destino
    gerar_arquivos(config)

def gerar_arquivos_por_percentual(quantidade_total, percentual_por_tipo, tipos_ativados=None, tamanhos_mb=None, diretorio_destino=None):
    """
    Gera arquivos com distribuição por percentual
    
    Args:
        quantidade_total: Quantidade total de arquivos
        percentual_por_tipo: Dicionário com percentuais (ex: {"pdf": 70, "outros": 30})
        tipos_ativados: Lista de tipos ativados
        tamanhos_mb: Tamanhos em MB por tipo
        diretorio_destino: Diretório de destino dos arquivos
    """
    config = ConfiguracaoArquivos()
    config.quantidade_total = quantidade_total
    config.percentual_por_tipo = percentual_por_tipo
    if tipos_ativados:
        config.tipos_ativados = tipos_ativados
    if tamanhos_mb:
        config.tamanho_mb.update(tamanhos_mb)
    if diretorio_destino:
        config.diretorio_destino = diretorio_destino
    gerar_arquivos(config)

def gerar_arquivos_por_template(quantidade_total, template="equilibrado", tipos_ativados=None, tamanhos_mb=None, diretorio_destino=None):
    """
    Gera arquivos usando templates de percentual pré-definidos do config.json
    
    Esta função facilita o uso de distribuições comuns sem precisar especificar
    percentuais manualmente. Usa templates pré-configurados no config.json.
    
    Args:
        quantidade_total (int): Quantidade total de arquivos a gerar
        template (str): Nome do template de percentual (padrão: "equilibrado")
        tipos_ativados (List[str], optional): Lista de tipos ativados
        tamanhos_mb (Dict[str, float], optional): Tamanhos em MB por tipo
        diretorio_destino (str, optional): Diretório de destino dos arquivos
        
    Templates Disponíveis:
        - "equilibrado": Distribuição igual entre todos os tipos (20% cada)
        - "foco_documentos": Foco em PDF (40%), DOCX (30%), TXT (20%), outros (10%)
        - "foco_dados": Foco em XLSX (50%), TXT (25%), PDF (15%), outros (10%)
        - "foco_imagens": Foco em JPEG (60%), PDF (20%), outros (20%)
        - "minimal": Apenas TXT (70%) e PDF (30%)
        
    Exemplo:
        >>> # Usar template equilibrado
        >>> gerar_arquivos_por_template(100, "equilibrado")
        
        >>> # Usar template focado em documentos
        >>> gerar_arquivos_por_template(50, "foco_documentos")
        
        >>> # Com diretório personalizado
        >>> gerar_arquivos_por_template(30, "foco_dados", diretorio_destino="meus_arquivos")
    """
    # Obter percentuais do template
    percentuais = obter_percentuais_padrao(template)
    
    # Usar a função existente de percentual
    gerar_arquivos_por_percentual(
        quantidade_total=quantidade_total,
        percentual_por_tipo=percentuais,
        tipos_ativados=tipos_ativados,
        tamanhos_mb=tamanhos_mb,
        diretorio_destino=diretorio_destino
    )

if __name__ == "__main__":
    # Exemplo de uso com configuração personalizada
    print("=== Gerando arquivos com configuração personalizada ===")
    
    # Configuração 1: Apenas TXT e PDF, tamanhos específicos
    config1 = ConfiguracaoArquivos(
        tipos_ativados=["txt", "pdf"],
        quantidade_por_tipo={"txt": 3, "pdf": 2},
        tamanho_mb={"txt": 0.2, "pdf": 0.5}
    )
    gerar_arquivos(config1)
    
    print("\n=== Gerando arquivos aleatoriamente ===")
    # Configuração 2: Modo aleatório
    gerar_arquivos_aleatorios(10, ["jpeg", "xlsx", "txt"])
    
    print("\n=== Gerando arquivos por quantidade total ===")
    # Configuração 3: Quantidade total específica
    gerar_arquivos_por_quantidade(15, ["txt", "pdf", "docx"])
    
    print("\n=== Gerando arquivos por percentual ===")
    # Configuração 4: Distribuição por percentual (70% PDF, 30% outros)
    gerar_arquivos_por_percentual(
        quantidade_total=20,
        percentual_por_tipo={"pdf": 70, "outros": 30},
        tipos_ativados=["txt", "pdf", "docx", "xlsx"],
        tamanhos_mb={"txt": 0.1, "pdf": 0.3, "docx": 0.2, "xlsx": 0.1}
    )
    
    print("\n=== Exemplo com diretório personalizado ===")
    # Configuração 5: Usando diretório personalizado
    config5 = ConfiguracaoArquivos(
        tipos_ativados=["txt", "pdf"],
        quantidade_por_tipo={"txt": 2, "pdf": 1},
        diretorio_destino="meus_arquivos_teste"
    )
    gerar_arquivos(config5)
    
    print("\n=== Exemplo usando função de conveniência com diretório ===")
    # Usando função de conveniência com diretório personalizado
    gerar_arquivos_aleatorios(5, ["txt", "jpeg"], "arquivos_personalizados")
